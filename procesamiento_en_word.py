##############################################################################################
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from procesamiento_datos import totales, no_mineras, mes_ano
from config import correlativas
import logging

logger = logging.getLogger(__name__)


def generar_docx(vars_from_totales:dict,vars_from_no_mineras: dict,vars_from_mes_ano:dict)-> Document:
    """
    Esta función genera el documento de Word con el resumen de las exportaciones.
    
    Parámetros:
    - vars_from_totales: Diccionario con las variables de totales.
    - vars_from_no_mineras: Diccionario con las variables de no mineras.
    - vars_from_mes_ano: Diccionario con las variables de mes y año.

    Retorna:

    - doc: Documento de Word con el resumen de las exportaciones.

    """
    
    #Variables globales para todo el informe:
    mes=vars_from_mes_ano["mes"]
    ano=vars_from_mes_ano["ano"]
    ano_ant=vars_from_mes_ano["ano_ant"]
    
    #Variables de totales:
    expt_act_tot = vars_from_totales["Total Export Act"]
    tagvar_tot = vars_from_totales["Tag Totales"]
    var_exp_tot = vars_from_totales["Varianza Totales"]
    expt_ant_tot = vars_from_totales["Exportaciones Totales Ant"]
    
    #Variables de no mineras:
    expt_act_tot_no_min = vars_from_totales["NME Export Act"]
    tagvar_nm_tot = vars_from_totales["Tag NME"]
    var_nm_tot = vars_from_totales["Varianza NME"]
    expt_ant_tot_no_min = vars_from_totales["NME Export Ant"]
    
    #Conteo de empresas:
    conteo_emp = vars_from_totales["Conteo"]
    
    #Destinos:
    tag_var_dest=vars_from_no_mineras["Destinos Tag"]
    variacion_destinos=vars_from_no_mineras["Destinos Varianza"]
    porcentaje_destinos=vars_from_no_mineras["Destinos Participacion"]
    agrupado_por_pais=vars_from_no_mineras["agrupado_por_pais"]
    datos_principales_exportadores=vars_from_no_mineras["datos_principales_exportadores"]
    exportado_10_principales=vars_from_no_mineras["Destinos Exportaciones"]
    
    #Empresas:
    analisis_empresas=vars_from_no_mineras["analisis_empresas"]
    var_empresas_resumen=vars_from_no_mineras["RS Varianza"]
    tag_var_emp=vars_from_no_mineras["RS Tag"]
    porcentaje_top10_emp=vars_from_no_mineras["RS Participacion"]
    valor_exp_top_10_emp=vars_from_no_mineras["RS Exportaciones"]

    #Subsectores:
    analisis_subsectores=vars_from_no_mineras["analisis_subsectores"]
    total_productos=vars_from_no_mineras["Productos Exportaciones"]
    variacion_productos=vars_from_no_mineras["Productos Varianza"]
    tag_var_productos=vars_from_no_mineras["Productos Tag"]

    #Departamentos:
    top_5_departamentos=vars_from_no_mineras["top_5_departamentos"]
    percentage_of_total= vars_from_no_mineras["Departamentos Participacion"]
    combined_percentage_variation=vars_from_no_mineras["Departamentos Varianza"]
    
    #Venezuela:
    results_venezuela=vars_from_no_mineras["results_venezuela"]
    growth_label_venezuela=vars_from_no_mineras["Venezuela Tag"]
    variation_venezuela=vars_from_no_mineras["Venezuela Variacion"]
    top_5_sectors_venezuela=vars_from_no_mineras["top_5_sectors_venezuela"]
    formatted_variations_companies=vars_from_no_mineras["formatted_variations_companies"]
    




##############################################Inicialización del documento############################################################################################
    
    doc = Document()

###################################################################################################################################################################
    
#0. Resumen Inicial
    
    #Titulo Principal
    doc.add_heading('Resumen de Exportaciones \n'f'Enero - {mes} de {ano}', level=0)
    doc.add_heading()  # Linea nueva en blanco para formato.

    # Se genera el título de la sección del resumen inicial 
    
    doc.add_heading(f'EXPORTACIONES Enero - {mes} de {ano} (DANE-DIAN)', 1)

    # Totales
    p = doc.add_paragraph()
    p.add_run('- ').bold = True
    p.add_run(f"Las cifras Enero - {mes} de {ano}, las exportaciones totales de Colombia fueron USD$ {format_to_millions(expt_act_tot)} millones, con un {tagvar_tot} del {var_exp_tot:.1f}% frente al mismo periodo de {vars_from_mes_ano['ano_ant']} USD$ {format_to_millions(expt_ant_tot)} millones.")

    # No mineras
    p = doc.add_paragraph()
    p.add_run('- ').bold = True
    p.add_run(f"Entre Enero - {mes} de {ano}, las exportaciones *no minero energéticas* de Colombia fueron USD$ {format_to_millions(expt_act_tot_no_min)} millones, con un {tagvar_nm_tot} del {var_nm_tot:.1f}% de las exportaciones del mismo periodo de {vars_from_mes_ano['ano_ant']} USD$ {format_to_millions(expt_ant_tot_no_min)} millones.")

    # Conteo Empresas
    p = doc.add_paragraph()
    p.add_run('- ').bold = True
    p.add_run(f"Entre Enero - {ano} de {ano}, un total de {conteo_emp} empresas exportaron productos no minero energéticos por montos superiores a USD 10,000 desde Colombia.")

    # Disclaimer
    p = doc.add_paragraph()
    p.add_run('- ').bold = True
    disclaimer = ("Vale la pena tener en cuenta que los datos que arroja el DANE-DIAN mes a mes: tienen dos meses de rezago, no incluyen las exportaciones desde Zona Franca y no tienen en cuenta servicios diferentes a Editorial (es decir que la cifra real no minero energética podría ser más alta).")
    p.add_run(disclaimer)

###################################################################################################################################################################

#1. Top 10 destinos de la exportación no minero energética y las empresas.
    
    nme_10_t=doc.add_heading(f'Top 10 destinos exportaciones no minero energéticas Enero-{mes} de {ano} (DANE-DIAN)', 1)
    nme_10_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Resumen inicial. 
    doc.add_paragraph(f"- Los 10 principales destinos de las exportaciones no minero energéticas del país suman total USD {format_to_millions(exportado_10_principales)} millones.")
    doc.add_paragraph(f"- Con un {tag_var_dest} del {variacion_destinos:.1f}% en nuestras exportaciones no minero energéticas frente al mismo período del año anterior.")
    doc.add_paragraph(f"- Estos mercados representan el {porcentaje_destinos:.1f}% de las exportaciones no minero energéticas de Colombia entre Enero - {mes} de {ano}.")
   
   # Procesamiento de los datos, se hace con un for sobre los valores de los diccionarios.
    for index, (country, value) in enumerate(agrupado_por_pais.items()):
        # Get the top 3 companies for the country
        top_companies = datos_principales_exportadores[country]["Principales exportadores"]
        companies_text = ", ".join([f"{company} (USD {format_to_millions(value)} millones)" for company, value in top_companies.items()])
        
        variance = datos_principales_exportadores[country]["Variación"]
        change_text = f"crecimiento de {variance:.1f}%" if variance >= 0 else f"decrecimiento de {-variance:.1f}%"

        # Añade un párrafo para cada país
        p = doc.add_paragraph()
        runner = p.add_run(f"{index + 1}. {country} : USD {format_to_millions(value)} millones, {change_text} frente a Enero – {mes} de {vars_from_mes_ano['ano_ant']}; principales exportadores: {companies_text}.")
        font = runner.font
        font.size = Pt(11)


###################################################################################################################################################################    

# 2. Top 10 empresas exportadoras no minero energéticas Enero -mes actual, año actual

    #Titulo Principal:
    t_emp=doc.add_heading(f'Top 10 empresas exportadoras no minero energéticas Enero -{vars_from_mes_ano["mes"]} de {vars_from_mes_ano["ano"]}', level=1)
    t_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #Resumen Inicial:
    doc.add_paragraph(f'- Las 10 principales empresas exportadoras no minero energéticas del país suman total USD {format_to_millions(valor_exp_top_10_emp)} millones.')
    doc.add_paragraph(f'- Se ve un: {tag_var_emp} de sus exportaciones en {var_empresas_resumen:.1f}% frente al mismo periodo del año {vars_from_mes_ano["ano_ant"]}.')
    doc.add_paragraph(f'- Concentran el {porcentaje_top10_emp:.1f}% de las exportaciones no minero energéticas de Colombia entre Enero – {vars_from_mes_ano["mes"]} {vars_from_mes_ano["ano"]}.')
    doc.add_paragraph()  # Linea nueva en blanco para formato. 
    
    # Se hace otro ciclo for que itera sobre los análisis de las empresas:
    for idx, company in enumerate(analisis_empresas, start=1):
        company_info = analisis_empresas[company]
        departamentos = ', '.join(company_info["Top Departamentos"].index)
        destinos = ', '.join(company_info["Top Destinos"].index)
        doc.add_paragraph(f'{idx}. {company} --> USD {format_to_millions(company_info["Total 2023 Exports"])} millones, {company_info["Tendencia"].lower()} del {company_info["Porcentaje"]:.1f}% frente a Enero – {vars_from_mes_ano["mes"]} {vars_from_mes_ano["ano_ant"]}; origen: {departamentos} ; destino: {destinos}.')

####################################################################################################################################################################

 # 3. Top 10 productos exportados no minero energéticos Enero -mes actual, año actual
    
    #Titulo Principal:
    t_prod=doc.add_heading(f'Top 10 productos exportados no minero energéticos Enero - {vars_from_mes_ano["mes"]} de 2023', 1)
    t_prod.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Resumen inicial:
    total_export_value = sum([details["Valor exportado Actual"] for details in analisis_subsectores.values()])#Bien    
    doc.add_paragraph(f"• Los 10 principales productos exportados no minero energéticos suman un total de USD {total_export_value/1000000:.1f} millones.")
    doc.add_paragraph(f"• Presentan un {tag_var_productos} del {variacion_productos:.1f}% frente a Enero - {mes} de {vars_from_mes_ano['ano_ant']}.")
    doc.add_paragraph(f"• Concentran el {total_productos / expt_act_tot_no_min * 100:.1f}% de las exportaciones no minero energéticas de Colombia entre Enero -{mes} de {ano}.")
    doc.add_paragraph()  # Linea nueva en blanco para formato. 

    # Otro ciclo for que itera sobre los análisis de los subsectores:
    for idx, (subsector, details) in enumerate(analisis_subsectores.items(), 1):
        
        # Se extraen las variables del diccionario de detalles:
        valor_exportado_actual = details["Valor exportado Actual"]
        variation = details["Variacion_sub"]
        tag = "crecimiento" if details["Variacion_sub"] > 0 else "decrecimiento"
        
        # Le damos formato a los orígenes:
        origins_str = ', '.join([f"{origin} (USD {value/1000000:.1f} millones)" for origin, value in details['USD from Top 3 Origins'].items()])

        # Se añade un párrafo para cada subsector:
        formatted_string = f"{idx}. {subsector}. USD {valor_exportado_actual/1000000:.1f} millones, {tag} del {variation:.1f}% frente a Enero -{mes} de {vars_from_mes_ano['ano_ant']}; origen principal: {origins_str}."
        doc.add_paragraph(formatted_string)

###################################################################################################################################################################
    
#4. Análisis por depatamentos
    
    #Titulo Principal:
    t_dept=doc.add_heading(f'Top 5 departamentos no-mineroenergéticos Enero-{vars_from_mes_ano["mes"]} de {vars_from_mes_ano["ano"]}', 1)
    t_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #Resumen inicial:
    doc.add_paragraph(f"• Los cinco principales departamentos exportadores no minero energéticos suman un total de USD {top_5_departamentos.loc['COMBINED', correlativas[9]]/1000000:.1f} millones.")
    doc.add_paragraph(f"• Presentan un: {'crecimiento' if combined_percentage_variation > 0 else 'decrecimiento'} sus exportaciones en un {abs(combined_percentage_variation):.1f} % frente a Enero – {mes} de {ano} {ano_ant}.")
    doc.add_paragraph(f"• Concentran el: {percentage_of_total:.1f} % de las exportaciones no minero energéticas de Colombia en Enero – {mes} {ano}.")

    # Se hace un ciclo for que itera sobre el dataframe que contiene la información de los departamentos:
    for idx, (depto, row) in enumerate(top_5_departamentos.drop('COMBINED').iterrows(), start=1):
        tendencia = "crecimiento" if row['Variance Percentage'] > 0 else "decrecimiento"
        doc.add_paragraph(f"{idx}. {depto}. USD {row[correlativas[9]]/1000000:.1f} millones, {tendencia} del {abs(row['Variance Percentage']):.1f}% frente a Enero - {mes} de {ano_ant}.") 

###################################################################################################################################################################

#6. Analisis de Venezuela:

    #Titulo Principal:
    t_vzla=doc.add_heading('*Venezuela*', level=1)
    t_vzla.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Resumen del crecimiento de las exportaciones a Venezuela:
    doc.add_paragraph(f"• Entre Enero – {mes} del presente año las exportaciones no mineras hacia Venezuela han {growth_label_venezuela} en {variation_venezuela:.1f}%.")

    # Top 5 sectores con mayores exportaciones a Venezuela:
    sectors_str = ', '.join(top_5_sectors_venezuela.index)
    doc.add_paragraph(f"• Los sectores con mayores exportaciones al mercado son: {sectors_str}.")

    # Resultados de la búsqueda de empresas:
    companies_str_list = [f"{company} ({variation})" for company, variation in formatted_variations_companies.items()]
    companies_str = ', '.join(companies_str_list)
    doc.add_paragraph(f"• Las empresas con mayores exportaciones son: {companies_str}.")

#########################################Finalización del documento############################################################################################

    return doc

########################################################################################################################################

#Generación del word: 

def format_to_millions(value: float) -> str:
    """
    Esta función formatea un valor numérico en millones con un decimal. Y los devuelve como un string.
    Parámetros:
    - value: Valor numérico a formatear.
    Retorna:
    - value_in_millions: Valor numérico formateado en millones con un decimal.
    """
    value_in_millions = value / 106
    return '{:,.1f}'.format(value_in_millions)
    

